import os
import json
import csv
import time
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from openai import OpenAI


# -----------------------------
# Settings
# -----------------------------
CONTRACTS_FOLDER = Path("contracts")
OUTPUTS_FOLDER = Path("outputs")

# Start with 2 for testing.
# After it works, change this to None to process all contracts.
TEST_LIMIT = None


# -----------------------------
# Prompt
# -----------------------------
SYSTEM_PROMPT = """
You are an AI contract extraction assistant.

Extract structured information from the contract.

Return clean JSON only.

Important rules:
- Do not invent information.
- If a field is missing, write "Not found".
- For every important field, include a short evidence quote from the contract.
- Keep summaries concise.
- Separate exact values from explanations.
- Use the same language as the contract when quoting evidence.
- Keep evidence quotes short.
- The output must be valid JSON.

Return this JSON structure:

{
  "contract_title": {
    "value": "",
    "evidence": ""
  },
  "contracting_parties": [
    {
      "party_name": "",
      "role": "",
      "evidence": ""
    }
  ],
  "effective_date": {
    "value": "",
    "type": "specific / conditional / not found",
    "evidence": ""
  },
  "expiration_date": {
    "value": "",
    "type": "specific / conditional / not found",
    "evidence": ""
  },
  "renewal_date": {
    "value": "",
    "type": "specific / conditional / not found",
    "evidence": ""
  },
  "governing_law": {
    "value": "",
    "evidence": ""
  },
  "jurisdiction": {
    "value": "",
    "evidence": ""
  },
  "payment_terms": {
    "summary": "",
    "payment_frequency": "",
    "late_payment_penalty": "",
    "evidence": ""
  },
  "key_obligations": [
    {
      "party": "",
      "obligation": "",
      "evidence": ""
    }
  ],
  "deliverables": [
    {
      "deliverable": "",
      "evidence": ""
    }
  ],
  "termination_clause_summary": {
    "summary": "",
    "evidence": ""
  },
  "confidentiality_clause_summary": {
    "summary": "",
    "evidence": ""
  },
  "missing_information": [],
  "unclear_clauses": []
}
"""


# -----------------------------
# Load environment
# -----------------------------
load_dotenv()

endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
api_key = os.getenv("AZURE_OPENAI_API_KEY")
deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT")

if not endpoint or not api_key or not deployment:
    raise ValueError(
        "Missing .env values. Check AZURE_OPENAI_ENDPOINT, "
        "AZURE_OPENAI_API_KEY, and AZURE_OPENAI_DEPLOYMENT."
    )

client = OpenAI(
    api_key=api_key,
    base_url=f"{endpoint.rstrip('/')}/openai/v1/"
)


# -----------------------------
# Helper functions
# -----------------------------
def read_docx(file_path: Path) -> str:
    """Extract readable text from a DOCX file, including tables."""
    document = Document(file_path)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_contract_data(contract_text: str) -> dict:
    """Send contract text to Azure OpenAI and return parsed JSON."""
    response = client.chat.completions.create(
        model=deployment,
        temperature=0,
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": SYSTEM_PROMPT
            },
            {
                "role": "user",
                "content": f"Contract text:\n{contract_text}"
            }
        ]
    )

    content = response.choices[0].message.content
    return json.loads(content)


def main():
    OUTPUTS_FOLDER.mkdir(exist_ok=True)

    docx_files = sorted(CONTRACTS_FOLDER.glob("*.docx"))

    if not docx_files:
        print("No DOCX files found in the contracts folder.")
        return

    if TEST_LIMIT is not None:
        docx_files = docx_files[:TEST_LIMIT]

    summary_rows = []

    for index, file_path in enumerate(docx_files, start=1):
        print(f"[{index}/{len(docx_files)}] Processing: {file_path.name}")

        try:
            contract_text = read_docx(file_path)

            if not contract_text.strip():
                raise ValueError("No text extracted from DOCX file.")

            extracted_data = extract_contract_data(contract_text)

            output_file = OUTPUTS_FOLDER / f"{file_path.stem}_extraction.json"

            with open(output_file, "w", encoding="utf-8") as f:
                json.dump(extracted_data, f, ensure_ascii=False, indent=2)

            summary_rows.append({
                "file": file_path.name,
                "status": "success",
                "output": output_file.name,
                "error": ""
            })

            print(f"Saved: {output_file}")

            # Small delay to avoid hitting rate limits too quickly
            time.sleep(1)

        except Exception as e:
            summary_rows.append({
                "file": file_path.name,
                "status": "failed",
                "output": "",
                "error": str(e)
            })

            print(f"Failed: {file_path.name}")
            print(f"Error: {e}")

    summary_file = OUTPUTS_FOLDER / "batch_summary.csv"

    with open(summary_file, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["file", "status", "output", "error"]
        )
        writer.writeheader()
        writer.writerows(summary_rows)

    print("\nBatch extraction finished.")
    print(f"Summary saved to: {summary_file}")


if __name__ == "__main__":
    main()