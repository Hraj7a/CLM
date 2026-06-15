import os
import json
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from openai import OpenAI
import streamlit as st
import pandas as pd


# -----------------------------
# Page config
# -----------------------------
st.set_page_config(
    page_title="AI Contract Analyzer",
    page_icon="📄",
    layout="wide"
)


# -----------------------------
# Load environment
# -----------------------------
load_dotenv()

endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
api_key = os.getenv("AZURE_OPENAI_API_KEY")
deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT")

if not endpoint or not api_key or not deployment:
    st.error("Missing Azure OpenAI settings. Check your .env file.")
    st.stop()

client = OpenAI(
    api_key=api_key,
    base_url=f"{endpoint.rstrip('/')}/openai/v1/"
)


# -----------------------------
# Prompts
# -----------------------------
EXTRACTION_PROMPT = """
You are an AI contract extraction assistant.

Extract structured information from the contract.

Return clean JSON only.

Important rules:
- Do not invent information.
- If a field is missing, write "Not found".
- For every important field, include a short evidence quote from the contract.
- Keep summaries concise.
- Separate exact values from explanations.
- Use the same language as the contract when quoting evidence.
- Keep evidence quotes short.
- The output must be valid JSON.

Return this JSON structure:

{
  "contract_title": {
    "value": "",
    "evidence": ""
  },
  "contracting_parties": [
    {
      "party_name": "",
      "role": "",
      "evidence": ""
    }
  ],
  "effective_date": {
    "value": "",
    "type": "specific / conditional / not found",
    "evidence": ""
  },
  "expiration_date": {
    "value": "",
    "type": "specific / conditional / not found",
    "evidence": ""
  },
  "renewal_date": {
    "value": "",
    "type": "specific / conditional / not found",
    "evidence": ""
  },
  "governing_law": {
    "value": "",
    "evidence": ""
  },
  "jurisdiction": {
    "value": "",
    "evidence": ""
  },
  "payment_terms": {
    "summary": "",
    "payment_frequency": "",
    "late_payment_penalty": "",
    "evidence": ""
  },
  "key_obligations": [
    {
      "party": "",
      "obligation": "",
      "evidence": ""
    }
  ],
  "deliverables": [
    {
      "deliverable": "",
      "evidence": ""
    }
  ],
  "termination_clause_summary": {
    "summary": "",
    "evidence": ""
  },
  "confidentiality_clause_summary": {
    "summary": "",
    "evidence": ""
  },
  "missing_information": [],
  "unclear_clauses": []
}
"""

RISK_PROMPT = """
You are an AI contract risk analysis assistant.

Your task is to analyze extracted contract data and identify legal, commercial, operational, and compliance risks.

Important rules:
- Do not invent facts.
- Base your analysis only on the provided extracted contract data.
- If information is missing, treat that as a possible risk.
- Do not provide final legal advice. Provide review points for a legal or compliance professional.
- Use clear and concise language.
- Use the same language as the extracted contract where appropriate.
- The output must be valid JSON only.

Risk categories to consider:
- Missing critical clauses
- Governing law and jurisdiction
- Payment and financial risk
- Termination risk
- Liability and indemnity risk
- Confidentiality risk
- Data protection / privacy risk
- Ambiguous obligations
- Renewal and expiry risk
- Operational compliance risk

Return this JSON structure:

{
  "contract_title": "",
  "overall_risk_level": "Low / Medium / High",
  "risk_score": 0,
  "risk_summary": "",
  "risks": [
    {
      "risk_title": "",
      "risk_level": "Low / Medium / High",
      "risk_category": "",
      "explanation": "",
      "evidence": "",
      "recommendation": ""
    }
  ],
  "missing_clauses": [
    {
      "clause_name": "",
      "risk_level": "Low / Medium / High",
      "why_it_matters": "",
      "recommendation": ""
    }
  ],
  "recommended_actions": [],
  "human_review_required": true
}

Scoring guide:
- 0 to 30 = Low risk
- 31 to 70 = Medium risk
- 71 to 100 = High risk
"""

QA_PROMPT = """
You are an AI contract question-answering assistant.

Your task is to answer user questions about a contract.

Important rules:
- Base your answer only on the provided contract text, extracted data, and risk analysis.
- Do not invent facts.
- If the answer is not found, say "Not found".
- Provide a short evidence quote where possible.
- Do not provide final legal advice.
- If the question asks for legal judgment, provide review points for a legal professional.
- The output must be valid JSON only.

Return this JSON structure:

{
  "answer": "",
  "evidence": "",
  "source": "contract text / extraction / risk analysis / not found",
  "confidence": "Low / Medium / High",
  "human_review_recommended": true
}
"""

# -----------------------------
# Helper functions
# -----------------------------
def read_docx(uploaded_file) -> str:
    document = Document(uploaded_file)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def call_model(system_prompt: str, user_content: str) -> dict:
    response = client.chat.completions.create(
        model=deployment,
        temperature=0,
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_content
            }
        ]
    )

    content = response.choices[0].message.content
    return json.loads(content)


def get_value(data: dict, field: str):
    item = data.get(field, {})
    if isinstance(item, dict):
        return item.get("value", "Not found")
    return "Not found"


def display_risk_badge(level: str):
    if level == "High":
        st.error("High Risk")
    elif level == "Medium":
        st.warning("Medium Risk")
    elif level == "Low":
        st.success("Low Risk")
    else:
        st.info(level or "Unknown")


# -----------------------------
# UI
# -----------------------------
st.title("AI-Driven Contract Lifecycle Management")
st.write("Upload a contract and generate extraction + risk analysis.")

uploaded_file = st.file_uploader(
    "Upload contract",
    type=["docx"]
)

if uploaded_file is None:
    st.info("Upload a DOCX contract to begin.")
    st.stop()

st.success(f"Uploaded: {uploaded_file.name}")

if st.button("Analyze Contract"):
    with st.spinner("Reading contract..."):
        contract_text = read_docx(uploaded_file)

    if not contract_text.strip():
        st.error("No text could be extracted from this document.")
        st.stop()

    with st.spinner("Extracting contract data..."):
        extraction_data = call_model(
            EXTRACTION_PROMPT,
            f"Contract text:\n{contract_text}"
        )

    with st.spinner("Analyzing contract risks..."):
        risk_data = call_model(
            RISK_PROMPT,
            "Analyze the following extracted contract data for risks:\n\n"
            + json.dumps(extraction_data, ensure_ascii=False, indent=2)
        )

    st.session_state["extraction_data"] = extraction_data
    st.session_state["risk_data"] = risk_data
    st.session_state["contract_text"] = contract_text
    st.session_state["uploaded_filename"] = uploaded_file.name


# -----------------------------
# Display results
# -----------------------------
if "extraction_data" not in st.session_state or "risk_data" not in st.session_state:
    st.stop()

extraction = st.session_state["extraction_data"]
risk = st.session_state["risk_data"]

st.markdown("---")
st.subheader("Contract Overview")

col1, col2, col3 = st.columns(3)

with col1:
    st.write("Contract Title")
    st.info(get_value(extraction, "contract_title"))

with col2:
    st.write("Risk Level")
    display_risk_badge(risk.get("overall_risk_level", "Unknown"))

with col3:
    st.metric("Risk Score", risk.get("risk_score", 0))


st.markdown("---")
st.subheader("Risk Summary")
st.write(risk.get("risk_summary", "No risk summary found."))


st.markdown("---")
st.subheader("Key Contract Information")

info_col1, info_col2 = st.columns(2)

with info_col1:
    st.write("Effective Date")
    st.info(get_value(extraction, "effective_date"))

    st.write("Expiration Date")
    st.info(get_value(extraction, "expiration_date"))

with info_col2:
    st.write("Governing Law")
    st.info(get_value(extraction, "governing_law"))

    st.write("Jurisdiction")
    st.info(get_value(extraction, "jurisdiction"))


st.markdown("---")
st.subheader("Contracting Parties")

parties = extraction.get("contracting_parties", [])

if parties:
    st.dataframe(pd.DataFrame(parties), use_container_width=True)
else:
    st.warning("No contracting parties found.")


st.markdown("---")
st.subheader("Payment Terms")

payment_terms = extraction.get("payment_terms", {})

if isinstance(payment_terms, dict):
    st.write("Summary")
    st.info(payment_terms.get("summary", "Not found"))

    st.write("Payment Frequency")
    st.info(payment_terms.get("payment_frequency", "Not found"))

    st.write("Late Payment Penalty")
    st.info(payment_terms.get("late_payment_penalty", "Not found"))
else:
    st.info("Payment terms not found.")


st.markdown("---")
st.subheader("Key Obligations")

obligations = extraction.get("key_obligations", [])

if obligations:
    st.dataframe(pd.DataFrame(obligations), use_container_width=True)
else:
    st.warning("No obligations found.")


st.markdown("---")
st.subheader("Top Risks")

risks = risk.get("risks", [])

if risks:
    for item in risks[:5]:
        with st.expander(f"{item.get('risk_level', '')} — {item.get('risk_title', '')}"):
            st.write("Category:")
            st.write(item.get("risk_category", "Not found"))

            st.write("Explanation:")
            st.write(item.get("explanation", "Not found"))

            st.write("Evidence:")
            st.write(item.get("evidence", "Not found"))

            st.write("Recommendation:")
            st.write(item.get("recommendation", "Not found"))
else:
    st.info("No risks found.")


st.markdown("---")
st.subheader("Missing Clauses")

missing_clauses = risk.get("missing_clauses", [])

if missing_clauses:
    st.dataframe(pd.DataFrame(missing_clauses), use_container_width=True)
else:
    st.success("No missing clauses identified.")


st.markdown("---")
st.subheader("Recommended Actions")

actions = risk.get("recommended_actions", [])

if actions:
    for action in actions:
        st.write(f"- {action}")
else:
    st.info("No recommended actions found.")


st.markdown("---")

with st.expander("View Extraction JSON"):
    st.json(extraction)

with st.expander("View Risk Analysis JSON"):
    st.json(risk)
# -----------------------------
# Contract Q&A
# -----------------------------
st.markdown("---")
st.subheader("Ask Questions About This Contract")

st.write("Ask a question in plain English or Arabic.")

question = st.text_input(
    "Your question",
    placeholder="Example: What are the payment terms?"
)

if st.button("Ask Question"):
    if not question.strip():
        st.warning("Please enter a question.")
    else:
        contract_text = st.session_state.get("contract_text", "")

        qa_input = {
            "user_question": question,
            "contract_text": contract_text,
            "extracted_data": extraction,
            "risk_analysis": risk
        }

        with st.spinner("Answering question..."):
            qa_result = call_model(
                QA_PROMPT,
                "Answer the user's question using this contract information:\n\n"
                + json.dumps(qa_input, ensure_ascii=False, indent=2)
            )

        st.session_state["last_qa_result"] = qa_result

if "last_qa_result" in st.session_state:
    qa_result = st.session_state["last_qa_result"]

    st.markdown("### Answer")
    st.info(qa_result.get("answer", "Not found"))

    st.markdown("### Evidence")
    st.write(qa_result.get("evidence", "Not found"))

    qa_col1, qa_col2, qa_col3 = st.columns(3)

    with qa_col1:
        st.metric("Confidence", qa_result.get("confidence", "Unknown"))

    with qa_col2:
        st.metric("Source", qa_result.get("source", "Unknown"))

    with qa_col3:
        review_needed = qa_result.get("human_review_recommended", True)
        st.metric("Human Review", "Yes" if review_needed else "No")