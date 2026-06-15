import os
import json

from dotenv import load_dotenv
from docx import Document
from openai import OpenAI

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware


# -----------------------------
# Load environment
# -----------------------------
load_dotenv()

endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
api_key = os.getenv("AZURE_OPENAI_API_KEY")
deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT")

if not endpoint or not api_key or not deployment:
    raise ValueError(
        "Missing .env values. Check AZURE_OPENAI_ENDPOINT, "
        "AZURE_OPENAI_API_KEY, and AZURE_OPENAI_DEPLOYMENT."
    )

client = OpenAI(
    api_key=api_key,
    base_url=f"{endpoint.rstrip('/')}/openai/v1/"
)


# -----------------------------
# FastAPI app
# -----------------------------
app = FastAPI(
    title="AI Contract Lifecycle Management API",
    version="1.0.0"
)

# Allows React frontend to call this backend during development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # For MVP only. Later restrict this.
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# -----------------------------
# Prompts
# -----------------------------
EXTRACTION_PROMPT = """
You are an AI contract extraction assistant.

Extract structured information from the contract.

Return clean JSON only.

Important rules:
- Do not invent information.
- If a field is missing, write "Not found".
- For every important field, include a short evidence quote from the contract.
- Keep summaries concise.
- Separate exact values from explanations.
- Use the same language as the contract when quoting evidence.
- Keep evidence quotes short.
- The output must be valid JSON.

Return this JSON structure:

{
  "contract_title": {
    "value": "",
    "evidence": ""
  },
  "contracting_parties": [
    {
      "party_name": "",
      "role": "",
      "evidence": ""
    }
  ],
  "effective_date": {
    "value": "",
    "type": "specific / conditional / not found",
    "evidence": ""
  },
  "expiration_date": {
    "value": "",
    "type": "specific / conditional / not found",
    "evidence": ""
  },
  "renewal_date": {
    "value": "",
    "type": "specific / conditional / not found",
    "evidence": ""
  },
  "governing_law": {
    "value": "",
    "evidence": ""
  },
  "jurisdiction": {
    "value": "",
    "evidence": ""
  },
  "payment_terms": {
    "summary": "",
    "payment_frequency": "",
    "late_payment_penalty": "",
    "evidence": ""
  },
  "key_obligations": [
    {
      "party": "",
      "obligation": "",
      "evidence": ""
    }
  ],
  "deliverables": [
    {
      "deliverable": "",
      "evidence": ""
    }
  ],
  "termination_clause_summary": {
    "summary": "",
    "evidence": ""
  },
  "confidentiality_clause_summary": {
    "summary": "",
    "evidence": ""
  },
  "missing_information": [],
  "unclear_clauses": []
}
"""

RISK_PROMPT = """
You are an AI contract risk analysis assistant.

Your task is to analyze extracted contract data and identify legal, commercial, operational, and compliance risks.

Important rules:
- Do not invent facts.
- Base your analysis only on the provided extracted contract data.
- If information is missing, treat that as a possible risk.
- Do not provide final legal advice. Provide review points for a legal or compliance professional.
- Use clear and concise language.
- Use the same language as the extracted contract where appropriate.
- The output must be valid JSON only.

Risk categories to consider:
- Missing critical clauses
- Governing law and jurisdiction
- Payment and financial risk
- Termination risk
- Liability and indemnity risk
- Confidentiality risk
- Data protection / privacy risk
- Ambiguous obligations
- Renewal and expiry risk
- Operational compliance risk

Return this JSON structure:

{
  "contract_title": "",
  "overall_risk_level": "Low / Medium / High",
  "risk_score": 0,
  "risk_summary": "",
  "risks": [
    {
      "risk_title": "",
      "risk_level": "Low / Medium / High",
      "risk_category": "",
      "explanation": "",
      "evidence": "",
      "recommendation": ""
    }
  ],
  "missing_clauses": [
    {
      "clause_name": "",
      "risk_level": "Low / Medium / High",
      "why_it_matters": "",
      "recommendation": ""
    }
  ],
  "recommended_actions": [],
  "human_review_required": true
}

Scoring guide:
- 0 to 30 = Low risk
- 31 to 70 = Medium risk
- 71 to 100 = High risk
"""

QA_PROMPT = """
You are an AI contract question-answering assistant.

Your task is to answer user questions about a contract.

Important rules:
- Base your answer only on the provided contract text, extracted data, and risk analysis.
- Do not invent facts.
- If the answer is not found, say "Not found".
- Provide a short evidence quote where possible.
- Do not provide final legal advice.
- If the question asks for legal judgment, provide review points for a legal professional.
- The output must be valid JSON only.

Return this JSON structure:

{
  "answer": "",
  "evidence": "",
  "source": "contract text / extraction / risk analysis / not found",
  "confidence": "Low / Medium / High",
  "human_review_recommended": true
}
"""


# -----------------------------
# Helper functions
# -----------------------------
def read_docx(uploaded_file) -> str:
    document = Document(uploaded_file.file)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def call_model(system_prompt: str, user_content: str) -> dict:
    response = client.chat.completions.create(
        model=deployment,
        temperature=0,
        response_format={"type": "json_object"},
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_content
            }
        ]
    )

    content = response.choices[0].message.content
    return json.loads(content)


def get_value(data: dict, field: str):
    item = data.get(field, {})
    if isinstance(item, dict):
        return item.get("value", "Not found")
    return "Not found"


def build_dashboard_summary(extraction: dict, risk: dict) -> dict:
    parties = extraction.get("contracting_parties", [])
    missing_info = extraction.get("missing_information", [])
    missing_clauses = risk.get("missing_clauses", [])
    risks = risk.get("risks", [])

    top_risks = []

    for item in risks[:5]:
        top_risks.append({
            "risk_title": item.get("risk_title", ""),
            "risk_level": item.get("risk_level", ""),
            "risk_category": item.get("risk_category", ""),
            "recommendation": item.get("recommendation", "")
        })

    return {
        "contract_title": get_value(extraction, "contract_title"),
        "overall_risk_level": risk.get("overall_risk_level", ""),
        "risk_score": risk.get("risk_score", 0),
        "risk_summary": risk.get("risk_summary", ""),
        "parties_count": len(parties) if isinstance(parties, list) else 0,
        "effective_date": get_value(extraction, "effective_date"),
        "expiration_date": get_value(extraction, "expiration_date"),
        "governing_law": get_value(extraction, "governing_law"),
        "jurisdiction": get_value(extraction, "jurisdiction"),
        "missing_information_count": len(missing_info) if isinstance(missing_info, list) else 0,
        "missing_clauses_count": len(missing_clauses) if isinstance(missing_clauses, list) else 0,
        "human_review_required": risk.get("human_review_required", True),
        "top_risks": top_risks
    }


# -----------------------------
# API routes
# -----------------------------
@app.get("/")
def root():
    return {
        "message": "AI Contract Lifecycle Management API is running"
    }


@app.post("/analyze")
async def analyze_contract(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail="Only DOCX files are supported for now."
        )

    try:
        contract_text = read_docx(file)

        if not contract_text.strip():
            raise HTTPException(
                status_code=400,
                detail="No readable text found in the document."
            )

        extraction = call_model(
            EXTRACTION_PROMPT,
            f"Contract text:\n{contract_text}"
        )

        risk = call_model(
            RISK_PROMPT,
            "Analyze the following extracted contract data for risks:\n\n"
            + json.dumps(extraction, ensure_ascii=False, indent=2)
        )

        dashboard_summary = build_dashboard_summary(extraction, risk)

        return {
            "filename": file.filename,
            "dashboard_summary": dashboard_summary,
            "extraction": extraction,
            "risk_analysis": risk,
            "contract_text": contract_text
        }

    except HTTPException:
        raise

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )


@app.post("/ask")
async def ask_contract_question(payload: dict):
    try:
        question = payload.get("question", "")
        contract_text = payload.get("contract_text", "")
        extraction = payload.get("extraction", {})
        risk_analysis = payload.get("risk_analysis", {})

        if not question.strip():
            raise HTTPException(
                status_code=400,
                detail="Question is required."
            )

        qa_input = {
            "user_question": question,
            "contract_text": contract_text,
            "extracted_data": extraction,
            "risk_analysis": risk_analysis
        }

        answer = call_model(
            QA_PROMPT,
            "Answer the user's question using this contract information:\n\n"
            + json.dumps(qa_input, ensure_ascii=False, indent=2)
        )

        return answer

    except HTTPException:
        raise

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )